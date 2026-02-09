"""
Tests for report formatters.
Verifies DOCX, PDF, and Markdown formatters produce correct output
with proper formatting, diagrams, and structure.
"""

import pytest
from pathlib import Path
import tempfile
import shutil
from src.formatters.markdown_formatter import MarkdownFormatter
from src.formatters.docx_formatter import DocxFormatter
from src.formatters.pdf_formatter import PdfFormatter


@pytest.fixture
def sample_sections():
    """Sample report sections for testing."""
    return {
        "abstract": "This is a **test abstract** with *italic* text and `code`.",
        "introduction": """## Introduction

This is the introduction section.

It has multiple paragraphs.

- Bullet point 1
- Bullet point 2
- Bullet point 3""",
        "methodology": """## Methodology

We used the following approach:

1. Step one
2. Step two
3. Step three

| Method | Accuracy | F1-Score |
|--------|----------|----------|
| Model A | 0.95 | 0.93 |
| Model B | 0.92 | 0.90 |
| Model C | 0.97 | 0.95 |
""",
        "results": """## Results

The results show excellent performance.

```python
def example():
    return "test"
```
""",
        "discussion": "This is the discussion section discussing the results.",
        "conclusion": "This is the conclusion section."
    }


@pytest.fixture
def sample_diagrams():
    """Sample Mermaid diagrams for testing."""
    return {
        "architecture": """graph TD
    A[User Input] --> B[Data Processing]
    B --> C[Model Training]
    C --> D[Predictions]
    D --> E[Output]""",
        "data_flow": """flowchart LR
    A[Raw Data] --> B[Preprocessing]
    B --> C[Feature Engineering]
    C --> D[Model]
    D --> E[Results]""",
        "metadata": {
            "format": "mermaid",
            "count": 2,
            "types": ["architecture", "data_flow"]
        }
    }


@pytest.fixture
def sample_citations():
    """Sample citations for testing."""
    return {
        "references_section": """# References

1. Smith, J. (2024). Machine Learning Basics. Journal of AI.
2. Jones, A. (2023). Data Science Fundamentals. Tech Press.
3. Brown, R. (2024). Advanced Analytics. Data Publishers.
"""
    }


@pytest.fixture
def test_config():
    """Test configuration."""
    return {
        "title": "Test Technical Report",
        "include_table_of_contents": True,
        "include_appendix": False
    }


@pytest.fixture
def temp_output_dir():
    """Create temporary directory for test outputs."""
    temp_dir = tempfile.mkdtemp()
    yield Path(temp_dir)
    # Cleanup after tests
    shutil.rmtree(temp_dir, ignore_errors=True)


class TestMarkdownFormatter:
    """Test cases for Markdown formatter."""
    
    def test_basic_formatting(self, test_config, sample_sections, sample_diagrams, 
                            sample_citations, temp_output_dir):
        """Test basic markdown generation."""
        formatter = MarkdownFormatter(test_config)
        output_path = temp_output_dir / "test_report.md"
        
        result = formatter.format_report(
            sections=sample_sections,
            diagrams=sample_diagrams,
            citations=sample_citations,
            output_path=output_path
        )
        
        assert result.exists()
        content = result.read_text()
        
        # Check title
        assert "# Test Technical Report" in content
        
        # Check sections
        assert "## Abstract" in content
        assert "## Introduction" in content
        assert "## Methodology" in content
        
        # Check diagrams
        assert "## Diagrams" in content
        assert "```mermaid" in content
        
        # Check references
        assert "# References" in content
    
    def test_table_of_contents(self, test_config, sample_sections, sample_diagrams,
                              sample_citations, temp_output_dir):
        """Test table of contents generation."""
        formatter = MarkdownFormatter(test_config)
        output_path = temp_output_dir / "test_toc.md"
        
        result = formatter.format_report(
            sections=sample_sections,
            diagrams=sample_diagrams,
            citations=sample_citations,
            output_path=output_path
        )
        
        content = result.read_text()
        assert "## Table of Contents" in content
        assert "[Introduction](#introduction)" in content
        assert "[Methodology](#methodology)" in content


class TestDocxFormatter:
    """Test cases for DOCX formatter."""
    
    def test_docx_creation(self, test_config, sample_sections, sample_diagrams,
                          sample_citations, temp_output_dir):
        """Test DOCX document creation."""
        try:
            formatter = DocxFormatter(test_config)
            output_path = temp_output_dir / "test_report.docx"
            
            result = formatter.format_report(
                sections=sample_sections,
                diagrams=sample_diagrams,
                citations=sample_citations,
                output_path=output_path
            )
            
            assert result.exists()
            assert result.suffix == '.docx'
            assert result.stat().st_size > 0
            
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_docx_with_tables(self, test_config, sample_sections, sample_diagrams,
                             sample_citations, temp_output_dir):
        """Test DOCX with markdown tables."""
        try:
            from docx import Document
            
            formatter = DocxFormatter(test_config)
            output_path = temp_output_dir / "test_tables.docx"
            
            result = formatter.format_report(
                sections=sample_sections,
                diagrams=sample_diagrams,
                citations=sample_citations,
                output_path=output_path
            )
            
            # Load and verify the document
            doc = Document(str(result))
            
            # Check that tables were created
            assert len(doc.tables) > 0
            
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_docx_heading_hierarchy(self, test_config, sample_sections, sample_diagrams,
                                   sample_citations, temp_output_dir):
        """Test proper heading hierarchy in DOCX."""
        try:
            from docx import Document
            
            formatter = DocxFormatter(test_config)
            output_path = temp_output_dir / "test_headings.docx"
            
            result = formatter.format_report(
                sections=sample_sections,
                diagrams=sample_diagrams,
                citations=sample_citations,
                output_path=output_path
            )
            
            doc = Document(str(result))
            
            # Verify headings exist
            headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
            assert len(headings) > 0
            
        except ImportError:
            pytest.skip("python-docx not available")


class TestPdfFormatter:
    """Test cases for PDF formatter."""
    
    def test_pdf_creation(self, test_config, sample_sections, sample_diagrams,
                         sample_citations, temp_output_dir):
        """Test PDF document creation."""
        try:
            formatter = PdfFormatter(test_config)
            output_path = temp_output_dir / "test_report.pdf"
            
            result = formatter.format_report(
                sections=sample_sections,
                diagrams=sample_diagrams,
                citations=sample_citations,
                output_path=output_path
            )
            
            assert result.exists()
            assert result.suffix == '.pdf'
            assert result.stat().st_size > 0
            
            # Verify PDF header
            with open(result, 'rb') as f:
                header = f.read(4)
                assert header == b'%PDF'
            
        except ImportError:
            pytest.skip("reportlab not available")
    
    def test_pdf_structure(self, test_config, sample_sections, sample_diagrams,
                          sample_citations, temp_output_dir):
        """Test PDF has proper structure."""
        try:
            formatter = PdfFormatter(test_config)
            output_path = temp_output_dir / "test_structure.pdf"
            
            result = formatter.format_report(
                sections=sample_sections,
                diagrams=sample_diagrams,
                citations=sample_citations,
                output_path=output_path
            )
            
            # Verify file size is reasonable (has content)
            assert result.stat().st_size > 5000  # At least 5KB
            
        except ImportError:
            pytest.skip("reportlab not available")


class TestFormatterIntegration:
    """Integration tests for all formatters."""
    
    def test_all_formatters_produce_output(self, test_config, sample_sections,
                                          sample_diagrams, sample_citations, temp_output_dir):
        """Test that all formatters can generate output."""
        formatters = [
            (MarkdownFormatter, "test.md"),
            (DocxFormatter, "test.docx"),
            (PdfFormatter, "test.pdf")
        ]
        
        for FormatterClass, filename in formatters:
            try:
                formatter = FormatterClass(test_config)
                output_path = temp_output_dir / filename
                
                result = formatter.format_report(
                    sections=sample_sections,
                    diagrams=sample_diagrams,
                    citations=sample_citations,
                    output_path=output_path
                )
                
                assert result.exists()
                assert result.stat().st_size > 0
                
            except ImportError:
                pytest.skip(f"{FormatterClass.__name__} dependencies not available")
    
    def test_empty_sections_handling(self, test_config, temp_output_dir):
        """Test formatters handle empty sections gracefully."""
        empty_sections = {}
        empty_diagrams = {"metadata": {"format": "mermaid", "count": 0, "types": []}}
        empty_citations = {}
        
        formatter = MarkdownFormatter(test_config)
        output_path = temp_output_dir / "empty_test.md"
        
        # Should not raise an error
        result = formatter.format_report(
            sections=empty_sections,
            diagrams=empty_diagrams,
            citations=empty_citations,
            output_path=output_path
        )
        
        assert result.exists()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
